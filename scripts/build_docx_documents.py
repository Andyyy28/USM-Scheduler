"""Generate polished Word versions of the concept paper and user guide."""

from __future__ import annotations

import argparse
import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

GREEN = "073A2A"
GREEN_MID = "125C40"
GOLD = "E0AD19"
GOLD_SOFT = "FFF4CC"
CANVAS = "F4F7F4"
INK = "16221B"
MUTED = "53635B"
RED_SOFT = "FFF1F0"

ROOT = Path(__file__).resolve().parents[1]
SEAL_PATH = ROOT / "scheduler" / "static" / "scheduler" / "img" / "usm-seal.png"

INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|<https?://[^>]+>|\*[^*]+\*)"
)
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
ORDERED_PATTERN = re.compile(r"^\s*\d+\.\s+(.+)$")
TABLE_SEPARATOR_PATTERN = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")


def _set_cell_fill(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:  # type: ignore[no-untyped-def]
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_shading(paragraph, fill: str) -> None:  # type: ignore[no-untyped-def]
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_paragraph_border(paragraph, *, color: str, side: str = "left") -> None:  # type: ignore[no-untyped-def]
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "16")
    border.set(qn("w:space"), "8")
    border.set(qn("w:color"), color)
    borders.append(border)


def _add_hyperlink(paragraph, text: str, url: str) -> None:  # type: ignore[no-untyped-def]
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN_MID)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_markdown(paragraph, text: str) -> None:  # type: ignore[no-untyped-def]
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(GREEN_MID)
            run.font.highlight_color = None
        elif token.startswith("["):
            label, url = re.match(r"^\[([^\]]+)\]\((https?://[^)]+)\)$", token).groups()  # type: ignore[union-attr]
            _add_hyperlink(paragraph, label, url)
        elif token.startswith("<"):
            _add_hyperlink(paragraph, token[1:-1], token[1:-1])
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    return re.sub(r"(?<!\*)\*([^*]+)\*", r"\1", text)


def _set_repeat_table_header(row) -> None:  # type: ignore[no-untyped-def]
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        if row_index == 0:
            _set_repeat_table_header(row)
        for column_index in range(column_count):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            value = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(_strip_inline_markdown(value))
            for run in paragraph.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.5)
            if row_index == 0:
                _set_cell_fill(cell, GREEN)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                _set_cell_fill(cell, CANVAS)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_special_start(lines: list[str], index: int) -> bool:
    line = lines[index]
    if not line.strip():
        return True
    return bool(
        HEADING_PATTERN.match(line)
        or line.startswith("```")
        or line.startswith(">")
        or line.lstrip().startswith("- ")
        or ORDERED_PATTERN.match(line)
        or (
            line.strip().startswith("|")
            and index + 1 < len(lines)
            and TABLE_SEPARATOR_PATTERN.match(lines[index + 1].strip())
        )
    )


def _add_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    start = next((index for index, line in enumerate(lines) if line.startswith("## ")), 0)
    index = start
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        heading = HEADING_PATTERN.match(line)
        if heading:
            level = min(max(len(heading.group(1)) - 1, 1), 3)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline_markdown(paragraph, heading.group(2))
            index += 1
            continue

        if line.startswith("```"):
            language = line[3:].strip()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            paragraph = document.add_paragraph(style="Code Block")
            _set_paragraph_shading(paragraph, CANVAS)
            _set_paragraph_border(paragraph, color=GREEN_MID)
            if language:
                label = paragraph.add_run(f"{language.upper()}\n")
                label.bold = True
                label.font.color.rgb = RGBColor.from_string(GREEN_MID)
            paragraph.add_run("\n".join(code_lines))
            continue

        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and TABLE_SEPARATOR_PATTERN.match(lines[index + 1].strip())
        ):
            table_rows = [_parse_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(_parse_table_row(lines[index]))
                index += 1
            _add_table(document, table_rows)
            continue

        if line.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].startswith(">"):
                quote_lines.append(lines[index][1:].strip())
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(5)
            paragraph.paragraph_format.right_indent = Mm(3)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            _set_paragraph_shading(paragraph, GOLD_SOFT)
            _set_paragraph_border(paragraph, color=GOLD)
            _add_inline_markdown(paragraph, " ".join(quote_lines))
            continue

        bullet = line.lstrip().startswith("- ")
        ordered = ORDERED_PATTERN.match(line)
        if bullet or ordered:
            content = line.lstrip()[2:] if bullet else ordered.group(1)  # type: ignore[union-attr]
            index += 1
            continuations: list[str] = []
            while index < len(lines) and lines[index].strip() and not _is_special_start(lines, index):
                continuations.append(lines[index].strip())
                index += 1
            if continuations:
                content = " ".join((content, *continuations))
            paragraph = document.add_paragraph(style="List Bullet" if bullet else "List Number")
            _add_inline_markdown(paragraph, content)
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not _is_special_start(lines, index):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = document.add_paragraph()
        _add_inline_markdown(paragraph, " ".join(paragraph_lines))


def _add_page_number(paragraph) -> None:  # type: ignore[no-untyped-def]
    paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _set_picture_alt_text(inline_shape, description: str) -> None:  # type: ignore[no-untyped-def]
    drawing = inline_shape._inline
    properties = drawing.docPr
    properties.set("descr", description)
    properties.set("title", "University of Southern Mindanao seal")


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    heading_sizes = {1: 15, 2: 12.5, 3: 11}
    for level, size in heading_sizes.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN if level == 1 else GREEN_MID)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(5)

    code_style = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(INK)
    code_style.paragraph_format.left_indent = Mm(5)
    code_style.paragraph_format.right_indent = Mm(3)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.keep_together = True
    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)


def _configure_document(document: Document, *, title: str, subject: str) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(21)
    section.bottom_margin = Mm(19)
    section.left_margin = Mm(23)
    section.right_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    properties = document.core_properties
    properties.title = title
    properties.subject = subject
    properties.author = "Ruby Jean B. Solomon; Edgardo Gabriel L. Paclibar"
    properties.keywords = "USM, university timetabling, CP-SAT, genetic algorithm, thesis prototype"
    properties.comments = "BSCS thesis prototype; not an official USM scheduling system."

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    header = section.header.paragraphs[0]
    header.text = "USM SCHEDULING DECISION SUPPORT · BSCS THESIS PROTOTYPE"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(GREEN_MID)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Not an official USM scheduling system  ·  ")
    _add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def _add_cover(
    document: Document,
    *,
    title: str,
    document_type: str,
    subtitle: str,
    metadata: Iterable[str],
) -> None:
    seal_paragraph = document.add_paragraph()
    seal_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = seal_paragraph.add_run().add_picture(str(SEAL_PATH), width=Inches(1.35))
    _set_picture_alt_text(shape, "Official University of Southern Mindanao seal")

    university = document.add_paragraph()
    university.alignment = WD_ALIGN_PARAGRAPH.CENTER
    university.paragraph_format.space_after = Pt(0)
    run = university.add_run("UNIVERSITY OF SOUTHERN MINDANAO")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(GREEN)

    location = document.add_paragraph("Kabacan, Cotabato")
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location.paragraph_format.space_after = Pt(18)
    for run in location.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    kind = document.add_paragraph(document_type.upper())
    kind.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kind.paragraph_format.space_after = Pt(9)
    for run in kind.runs:
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(GOLD)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(10)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = RGBColor.from_string(GREEN)

    subtitle_paragraph = document.add_paragraph(subtitle)
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(20)
    for run in subtitle_paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    for value in metadata:
        paragraph = document.add_paragraph(value)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)

    disclaimer = document.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(20)
    disclaimer.paragraph_format.space_after = Pt(0)
    disclaimer.paragraph_format.left_indent = Mm(15)
    disclaimer.paragraph_format.right_indent = Mm(15)
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_shading(disclaimer, RED_SOFT)
    _set_paragraph_border(disclaimer, color=GOLD, side="top")
    run = disclaimer.add_run(
        "BSCS THESIS PROTOTYPE — NOT AN OFFICIAL USM SCHEDULING SYSTEM. "
        "Institutional data, policies, approval, and seal use require formal authorization."
    )
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GREEN)

    document.add_page_break()


def _add_toc(document: Document) -> None:
    heading = document.add_paragraph("Contents", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    paragraph = document.add_paragraph()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Open in Microsoft Word to update the table of contents."
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend((field_begin, instruction, field_separate, placeholder, field_end))
    paragraph._p.append(run)
    document.add_page_break()


def build_document(
    markdown_path: Path,
    output_path: Path,
    *,
    title: str,
    document_type: str,
    subtitle: str,
    metadata: Iterable[str],
) -> None:
    markdown = markdown_path.read_text(encoding="utf-8")
    document = Document()
    _configure_styles(document)
    _configure_document(document, title=title, subject=document_type)
    _add_cover(
        document,
        title=title,
        document_type=document_type,
        subtitle=subtitle,
        metadata=metadata,
    )
    _add_toc(document)
    _add_markdown(document, markdown)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=ROOT / "docs")
    args = parser.parse_args()

    concept_title = (
        "A Comparative Evaluation of CP-SAT and Genetic Algorithm for University "
        "Timetabling and College-Boundary-Aware Room Assignment at the University "
        "of Southern Mindanao"
    )
    build_document(
        ROOT / "docs" / "concept-paper.md",
        args.output_dir / "USM-Scheduler-Concept-Paper.docx",
        title=concept_title,
        document_type="Concept Paper",
        subtitle="Kabacan Main Campus · One-campus, one-term case-study proposal",
        metadata=(
            "Ruby Jean B. Solomon",
            "Edgardo Gabriel L. Paclibar",
            "Bachelor of Science in Computer Science",
            "August 2026",
        ),
    )
    build_document(
        ROOT / "docs" / "user-guide.md",
        args.output_dir / "USM-Scheduler-User-Guide.docx",
        title="USM Scheduler User Guide",
        document_type="User Guide",
        subtitle="Professional operating guide for the scheduling decision-support thesis prototype",
        metadata=(
            "For system administrators, central scheduling personnel, and college reviewers",
            "Prepared by Ruby Jean B. Solomon and Edgardo Gabriel L. Paclibar",
            "Version 1.0 · August 2026",
        ),
    )
    print(f"Created Word documents in {args.output_dir.resolve()}")


if __name__ == "__main__":
    main()
